import docx
import re
from bs4 import BeautifulSoup
import requests
import webbrowser
import time


class Pelicula(object):
    def __init__(self, urlFA=None):
        self.url_FA = str(urlFA)
        resp = Pelicula.SafeGetUrl(self.url_FA)
        if resp.status_code == 404:
            self.exists = False
            return  # Si el id no es correcto, dejo de construir la clase
        else:
            self.exists = True

        # Parseo la página
        self.parsed_page = BeautifulSoup(resp.text,'html.parser')

        self.director = None
        self.año = None
        self.duración = None

    def GetDirectorYearDuration(self):

        l = self.parsed_page.find(id="left-column")
        try:
            self.duración = l.find(itemprop="duration").contents[0]
        except:
            return False
            # caso en el que no está escrita la duración
        # quito el sufijo min.
        self.duración = int(self.duración.split(' ', 1)[0])

        try:
            self.director = l.find(itemprop="director").contents[0].contents[0].contents[0]
        except:
            return False

        try:
            self.año = l.find(itemprop="datePublished").contents[0]
        except:
            return False

    @staticmethod
    def SafeGetUrl(url):
        # open with GET method
        resp = requests.get(url)
        # Caso 429: too many requests
        if resp.status_code == 429:
            return Pelicula.PassCaptcha(url)
        else:  # No está contemplado el caso 404: not found
            return resp

    @staticmethod
    def PassCaptcha(url):
        # abro un navegador para poder pasar el Captcha
        webbrowser.open(url)
        resp = requests.get(url)
        print("\nPor favor, entra en FilmAffinity y pasa el captcha por mí.")
        # Controlo que se haya pasado el Captcha
        while resp.status_code != 200:
            time.sleep(3)  # intento recargar la página cada 3 segundos
            resp = requests.get(url)
        return resp


class html():

    def __init__(self):
        # Abro el documento para leerlo
        self.doc = docx.Document('c:\\Users\\usuario\\Desktop\\Jorges things\\Pelécolas\\Reseñas.docx')

        self.parrafos_critica = []
        self.titulo = ""
        self.año = ""
        self.duración = ""

        self.titulos = {}  # Hago un diccionario con todos los títulos que tienen una crítica escrita
        search_title = False

        for i, paragraph in enumerate(self.doc.paragraphs): # Recorro todos los párrafos del documento
            if not search_title:
                if paragraph.text == 'Películas':
                    # El encabezado de las críticas cinematográficas no me interesa
                    continue
                elif paragraph.text == 'Literatura':
                    # Cuando llegue al apartado de literatura dejo de contar
                    break
                if self.__fin_de_parrafo(paragraph.text):
                    # Si hay un doble salto de párrafo es que ha terminado una crítica
                    search_title = True # El inicio del siguiente párrafo será el título de la película
            else:
                # Sé que estoy en un párrafo que es el primero de una crítica
                # Este párrafo comenzará con el título de la película
                # El título se separa de la crítica con dos puntos
                titulo = ""
                for run in paragraph.runs:
                # Conservo las negritas
                    if not run.bold:  # he llegado al final del título
                        break
                    titulo += run.text

                # Quito el separador, previsiblemente los dos puntos.
                titulo = titulo.strip(".: ")
                if not self.__fin_de_parrafo(titulo):
                    # Guardo el parrafo donde empieza la crítica.
                    self.titulos[titulo] = i
                else:
                    # Si hemos encontrado un doble salto de linea, hemos llegado a los libros
                    break
                search_title = False # Devuelvo la variable a su valor original

        self.__unwanted_chars = dict.fromkeys(map(ord, " ,!¡@#$?¿()."), None)
        self.__unwanted_chars.update(zip(map(ord, "áéíóú"),map(ord, "aeiou")))
        self.__unwanted_chars.update(zip(map(ord, "àèìòù"),map(ord, "aeiou")))
        self.__unwanted_chars.update(zip(map(ord, "âêîôû"),map(ord, "aeiou")))

    def ask_for_data(self):
        exists_given_title = False
        # Pido los datos de la película que voy a buscar
        while not exists_given_title:
            self.titulo = input("Introduzca título de la película: ")
            exists_given_title = self.exists(self.titulo)
        self.director = input("Introduzca director: ")
        if not self.__interpretate_director():
            return
        self.año = input("Introduzca el año: ")
        self.duración = input("Introduzca duración de la película: ")

    def __interpretate_director(self):
        # Si es un número, considero que se ha introducido un id de Filmaffinitty
        if self.director.isnumeric():
            url = 'https://www.filmaffinity.com/es/film' + self.director + '.html'
            return not self.__get_data_from_FA(url)
        if self.director.find("filmaffinity") >= 0:
            # Se ha introducido directamente la url
            return not self.__get_data_from_FA(self.director)
        else:
            # El director de la película es lo introducido por teclado
            return True

    def __get_data_from_FA(self, url):
        peli = Pelicula(urlFA=url)
        if not peli.exists:
            return False
        peli.GetDirectorYearDuration()
        self.director = peli.director
        self.año = peli.año
        self.duración = peli.duración
        return True

    def exists(self,titulo):
        # No quiero que sea sensible a las mayúsculas
        titulo = titulo.lower()

        for val in self.titulos:
            if titulo == val.lower():
                # Caso de coincidencia exacta
                self.titulo = val
                return True

        # Si es posible, siguiero los títulos más cercanos al introducido
        self.__closest_title(titulo)

        return False

    def __closest_title(self, titulo):
        # Intento buscar los casos más cercanos
        titulo = self.__normalize_string(titulo)
        found = False
        for iter in self.titulos:
            iter_ = self.__normalize_string(iter)
            if titulo.find(iter_) >= 0 or iter_.find(titulo) >= 0:
                # Hay suficiente concordancia como para sugerir el título
                if not found:
                    # Aún no he impreso nada por pantalla
                    print("Quizás quisiste decir...")
                    found = True
                # Imprimo el título original. El que se ha leído en el documento
                print(iter)

    def __normalize_string(self, str):
        # Elimino las mayúsculas
        str = str.lower()
        # Elimino espacios, signos de puntuación y acentos
        str = str.translate(self.__unwanted_chars)
        # Elimino caracteres repetidos
        str = re.sub(r'(.)\1+', r'\1', str)
        return str

    @staticmethod
    def __fin_de_parrafo(text):
        return text == "" or text == "\t"

    def search_film(self):
        # Si no tengo los datos de la película, los pido
        if not self.titulo:
            self.ask_for_data()
        # Empiezo a recorrer los párrafos desde el que sé que inicia la crítica que busco
        for paragraph in self.doc.paragraphs[self.titulos[self.titulo]:]:

            if self.__fin_de_parrafo(paragraph.text):
                # He llagado al final de la crítica. Dejo de leer el documento
                return self.parrafos_critica
            # Inicializo el párrafo
            parr_text = ""

            for run in paragraph.runs:
                # Conservo las cursivas
                if run.italic:
                    parr_text += "<i>" + run.text + "</i>"
                else:
                    parr_text += run.text

            if not self.parrafos_critica:
                # Si es el primer párrafo elimino el título
                parr_text = parr_text[len(self.titulo):]
                parr_text = parr_text.lstrip(": ")

            # Añado saltos de línea para un html más legible
            parr_text = parr_text.replace(". ", ".\n")
            # Añado el texto a la lista de párrafos
            self.parrafos_critica.append(parr_text)

        assert("No ha sido posible encontrar la reseña.")

    def write_html(self):

        self.search_film()

        reseña = open("Reseña " + str(self.titulo) + ".html", mode="w",encoding="utf-8")

        # Escribo el encabezado
        reseña.write("<!-- Encabezado -->\n")
        reseña.write("<div style=\"text-align: right;\">\n")
        reseña.write("<span style=\"font-family: &quot;courier new&quot; , &quot;courier&quot; , monospace;\">Dir.: " + str(self.director) + "</span></div>\n")
        reseña.write("<div style=\"text-align: right;\">\n")
        reseña.write("<span style=\"font-family: &quot;courier new&quot; , &quot;courier&quot; , monospace;\">" + str(self.año) + "</span></div>\n")
        reseña.write("<div style=\"text-align: right;\">\n")
        reseña.write("<span style=\"font-family: &quot;courier new&quot; , &quot;courier&quot; , monospace;\">" + str(self.duración) + " min.</span></div>\n")

        # Iteramos los párrafos
        reseña.write("\n<!-- Párrafos -->\n")
        for parrafo in self.parrafos_critica:
            reseña.write("<div style=\"margin: 16px 0px; text-align: justify; text-indent: 21.25pt;\">\n")
            reseña.write("<span style=\"font-family: &quot;times new roman&quot; , serif; margin: 0px;\">\n")
            reseña.write(str(parrafo))
            reseña.write("\n</span></div>\n")

        # Escribo los botones de Twitter
        reseña.write("\n<!--Boton follow-->\n")
        reseña.write("<a href=\"https://twitter.com/pelecolas?ref_src=twsrc%5Etfw\" class=\"twitter-follow-button\" data-show-count=\"false\">\n")
        reseña.write("Follow @pelecolas</a><script async src=\"https://platform.twitter.com/widgets.js\" charset=\"utf-8\"></script>\n")
        reseña.write("\n<!--Boton compartir-->\n")
        reseña.write("<a href=\"https://twitter.com/share?ref_src=twsrc%5Etfw\" class=\"twitter-share-button\" data-show-count=\"false\">Tweet</a><script async src=\"https://platform.twitter.com/widgets.js\" charset=\"utf-8\"></script>\n")

        reseña.close()

if __name__ == "__main__":
    Documento = html()
    Documento.write_html()
