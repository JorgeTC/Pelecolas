import docx
from dlg_config import CONFIG

SEPARATOR_YEAR = " - "

class WordReader():
    def __init__(self, folder):
        # Abro el documento para leerlo
        sz_doc_path = self.__get_word_file_name(folder)
        # Me quedo con el nombre del archivo sin la extensión.
        self.header = str(sz_doc_path[0].stem).split(SEPARATOR_YEAR)[0]
        self.folder = folder
        # Me guardo sólo los párrafos, es lo que voy a iterar más adelante
        self.paragraphs = []
        # Guardo a qué párrafo corresponde cada año
        self.years_parr = {}
        # Itero todos los docx que he encontrado
        for word in sz_doc_path:
            # Obtengo el año actual
            year = word.stem.split(SEPARATOR_YEAR)[1]
            # Guardo en qué párrafo empieza el año actual
            self.years_parr[year] = len(self.paragraphs)
            # Añado los párrafos del docx actual
            # Evito añadir el primero, donde está el título del documento
            self.paragraphs = self.paragraphs + docx.Document(word).paragraphs[1:]

        # Lista con todos los títulos que encuentre.
        self.titulos = {}

    def __get_title(self, paragraph):

        # Obtengo lo que esté en negrita
        title = self.__get_bold_title(paragraph)

        # Compruebo que lo que he leído sea un título.
        # Busco que su separador sean dos puntos.
        # Elimino los espacios que fueren delante del separador.
        title = title.strip(" ")
        # Compruebo que esté presente el separador.
        if not title or title[-1] != ":":
            title = ""
        else:
            # Ya he usado los dos puntos para reconocer el título.
            # Ahora los puedo eliminar para tener limpio el título.
            title = title.strip(": ")

        return title

    def __has_next_parr_title(self, text):
        if self.__is_header(text):
            # No cuento el encabezado del documento
            return False
        if self.__is_break_line(text):
            # Si hay un doble salto de párrafo, quizás ha terminado una crítica
            # El inicio del siguiente párrafo será el título de la película
            return True

        return False

    def __is_break_line(self, text):
        if text == '':
            return True
        if text == "\t":
            return True
        if text == "\n":
            return True

        return False

    def __is_header(self, text):
        # Quiero que el código valga para contar las películas y los libros.
        # No sé qué encabezado me voy a encontrar en mi documento.
        return text == self.header

    def __get_word_file_name(self, folder):
        # Carpeta donde están guardados los archivos word
        word_folder = CONFIG.get_value(CONFIG.S_COUNT_FILMS, CONFIG.P_WORD_FOLDER)
        # Si en el archivo de configuración se especifica una carpeta, busco en ella
        if word_folder:
            folder = folder / word_folder

        # Me espero un único archivo docx
        all_files = [x for x in folder.iterdir()]
        all_files = [x for x in all_files if x.suffix.lower() == ".docx"]

        return all_files

    def __append_title(self, paragraph, index):
        # Leo el posible título de este párrafo.
        titulo = self.__get_title(paragraph)
        # Si no se ha encontrado título, no es el inicio de una crítica.
        if not titulo:
            # No he conseguido añadir nada.
            return False

        # En este punto ya sabemos que tenemos un titulo
        self.titulos[titulo] = index

        # He añadido un título.
        return True

    def __get_bold_title(self, paragraph):
        # Obtengo el primer fragamento de texto que esté en negrita.
        titulo = ""

        for run in paragraph.runs:
            # Conservo las negritas
            if not run.bold:
                # he llegado al final del título
                break
            titulo += run.text

        return titulo

    def list_titles(self):
        # inicializo la variable.
        # No quiero buscar desde el principio porque sé que encontraré el título del documento.
        search_title = False

        # Recorro todos los párrafos del documento
        for i, paragraph in enumerate(self.paragraphs):
            if not search_title:
                # Compruebo si el próximo párrafo podría tener un título.
                search_title = self.__has_next_parr_title(paragraph.text)
            else:
                # Probablemente estoy en un párrafo que es el primero de una crítica.
                # Si he añadido un título, no me espero que el siguiente párrafo comience con título.
                # Si no he añadido nada, sigo buscando.
                search_title = not self.__append_title(paragraph, i)

        return self.titulos.keys()

    def write_list(self):
        # Abro el documento txt para escribirlo
        titulos_doc = open(self.folder / "Titulos de reseñas.txt", "w")

        # Miro si hay que escribir el índice
        b_index = CONFIG.get_bool(CONFIG.S_COUNT_FILMS, CONFIG.P_ADD_INDEX)

        # Miro si hay que escribir el año
        b_year = CONFIG.get_bool(CONFIG.S_COUNT_FILMS, CONFIG.P_ADD_YEAR)
        next_years = iter(self.years_parr.keys())
        # Cojo el primero de los años que hay que iterar
        next_year = next(next_years, None)

        for index, titulo in enumerate(self.titulos):

            # Escritura del año
            if b_year:
                # Compruebo que el título actual esté en una posición superior a donde inicia el siguiente año.
                # No me espero el caso de igualdad ya que ese párrafo corresponde al título del Word: peliculas.
                if self.titulos[titulo] > self.years_parr[next_year]:
                    # Escribo el año en el documento
                    titulos_doc.write("***" + str(next_year) + "***\n")
                    # Avanzo al siguiente año
                    past_year = next_year
                    next_year = next(next_years, past_year)
                    # Si es el último año, dejo de escribir años
                    if next_year == past_year:
                        b_year = False

            # Si tengo que añadir el índice cojo el número y añado un espacio
            if b_index:
                titulos_doc.write(str(index + 1) + " " + titulo + "\n")
            else:
                titulos_doc.write(titulo + "\n")

        # cierro el documento
        titulos_doc.close()
