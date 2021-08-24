import docx
import os

class WordReader():
    def __init__(self):
        # Abro el documento para leerlo
        sz_literatura = self.__get_word_file_name()
        # Me quedo con el nombre del archivo sin la extensión.
        self.header = sz_literatura[:sz_literatura.find(".")]
        # Me guardo sólo los párrafos, es lo que voy a iterar más adelante
        self.paragraphs = docx.Document(sz_literatura).paragraphs

        # Lista con todos los títulos que encuentre.
        self.titulos = []

    def __get_title(self, paragraph):

        # Obtengo lo que esté en negrita
        title = self.__get_bold_title(paragraph)

        # Compruebo que lo que he leído sea un título.
        # Busco que su separador sean dos puntos.
        # Elimino los espacios que fueren delante del separador.
        title = title.strip(" ")
        # Compruebo que esté presente el separador.
        if not title or title[-1] != ":":
            title = ""
        else:
            # Ya he usado los dos puntos para reconocer el título.
            # Ahora los puedo eliminar para tener limpio el título.
            title = title.strip(": ")

        return title

    def __has_next_parr_title(self, text):
        if self.__is_header(text):
            # No cuento el encabezado del documento
            return False
        if self.__is_break_line(text):
            # Si hay un doble salto de párrafo, quizás ha terminado una crítica
            # El inicio del siguiente párrafo será el título de la película
            return True

        return False

    def __is_break_line(self, text):
        if text == '':
            return True
        if text == "\t":
            return True
        if text == "\n":
            return True

        return False

    def __is_header(self, text):
        # Quiero que el código valga para contar las películas y los libros.
        # No sé qué encabezado me voy a encontrar en mi documento.
        return text == self.header

    def __get_word_file_name(self):
        # Me espero un único archivo docx
        all_files = os.listdir()
        all_files = [x for x in all_files if x.endswith(".docx")]

        return all_files[0]

    def __append_title(self, paragraph):
        # Leo el posible título de este párrafo.
        titulo = self.__get_title(paragraph)
        # Si no se ha encontrado título, no es el inicio de una crítica.
        if not titulo:
            # No he conseguido añadir nada.
            return False

        # En este punto ya sabemos que tenemos un titulo
        self.titulos.append(titulo)

        # He añadido un título.
        return True

    def __get_bold_title(self, paragraph):
        # Obtengo el primer fragamento de texto que esté en negrita.
        titulo = ""

        for run in paragraph.runs:
            # Conservo las negritas
            if not run.bold:
                # he llegado al final del título
                break
            titulo += run.text

        return titulo

    def list_titles(self):
        # inicializo la variable.
        # No quiero buscar desde el principio porque sé que encontraré el título del documento.
        search_title = False

        # Recorro todos los párrafos del documento
        for paragraph in self.paragraphs:
            if not search_title:
                # Compruebo si el próximo párrafo podría tener un título.
                search_title = self.__has_next_parr_title(paragraph.text)
            else:
                # Probablemente estoy en un párrafo que es el primero de una crítica.
                # Si he añadido un título, no me espero que el siguiente párrafo comience con título.
                # Si no he añadido nada, sigo buscando.
                search_title = not self.__append_title(paragraph)

        return self.titulos

    def write_list(self):
        titulos_doc = open("Titulos de reseñas.txt", "w")

        for titulo in self.titulos:
            titulos_doc.write(titulo + "\n")

        # cierro el documento
        titulos_doc.close()


if __name__ == "__main__":

    reader = WordReader()
    reader.list_titles()
    reader.write_list()

    print(len(reader.titulos))
    input("Press Enter to continue...")
